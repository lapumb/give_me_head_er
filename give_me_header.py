"""give_me_header.py

Replace or add an image in a .docx document header, where the header contains a table with at least one cell.

Usage:
    give_me_header.py --docx=<docx> --image=<image> [--width_inches=<width_inches>] [--height_inches=<height_inches>]
    give_me_header.py (-h | --help)
"""

from typing import Any, Optional
from docopt import docopt
import docx

def add_or_replace_image_in_header_table(docx_doc_path: str, new_image_path: str, width_inches: Optional[int] = 0, height_inches: Optional[int] = 0) -> None:
    if not docx_doc_path:
        print("No docx file given.")
        print("To see the usage, run the script with the --help option.")
        return
    
    if not new_image_path:
        print("No image file given.")
        print("To see the usage, run the script with the --help option.")
        return
    
    # Try to get the specified docx file
    try:
        docx_doc: docx.Document = docx.Document(docx_doc_path)
    except FileNotFoundError:
        print("Could not find the specified docx file.")
        return

    # Try to get the document header and the table in the header
    try:
        header: docx.section._Header = docx_doc.sections[0].header
        header_table: docx.table.Table = header.tables[0]
    except IndexError:
        print("Invalid header.")
        return
    
    # Clear the [0, 0] cell
    first_cell = header_table.rows[0].cells[0]
    first_cell._element.clear_content()
    
    # Add a "run" to the cell with a purpose to add an image
    cell_run: docx.text.run.Run = first_cell.add_paragraph().add_run()
    
    if width_inches and height_inches:
        cell_run.add_picture(new_image_path, width=docx.shared.Inches(width_inches), height=docx.shared.Inches(height_inches))
    elif width_inches:
        cell_run.add_picture(new_image_path, width=docx.shared.Inches(width_inches))
    elif height_inches:
        cell_run.add_picture(new_image_path, height=docx.shared.Inches(height_inches))
    else:
        cell_run.add_picture(new_image_path)
    
    # Save the document
    docx_doc.save(docx_doc_path)
        
if __name__ == '__main__':
    args: dict[str, Any] = docopt(__doc__)
    
    docx_doc_path: str = args['--docx']
    new_image_path: str = args['--image']
    width_inches: int = int(args['--width_inches']) if args['--width_inches'] else 0
    height_inches: str = int(args['--height_inches']) if args['--height_inches'] else 0
    
    add_or_replace_image_in_header_table(docx_doc_path, new_image_path, width_inches=width_inches, height_inches=height_inches)
